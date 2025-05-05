import os
import json
import hashlib
import shutil
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain.docstore.document import Document
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
from langchain.prompts.prompt import PromptTemplate
from langchain.embeddings.base import Embeddings
import docx
import pdfplumber
import chardet
from razdel import sentenize
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import warnings
import gradio as gr


# Загрузка переменных окружения
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    warnings.warn("OPENAI_API_KEY не найден в .env файле")

DATA_FOLDER = os.getenv("DATA_FOLDER", "legal_data")
INDEXES_FOLDER = os.getenv("INDEXES_FOLDER", "legal_indexes")
os.makedirs(DATA_FOLDER, exist_ok=True)
os.makedirs(INDEXES_FOLDER, exist_ok=True)

# Класс для создания эмбеддингов текста
class MyEmbeddings(Embeddings):
    def __init__(self):
        self.model = SentenceTransformer('all-MiniLM-L12-v2')

    def embed_documents(self, texts):
        return self.model.encode(texts, convert_to_numpy=True).tolist()

    def embed_query(self, text):
        return self.model.encode([text], convert_to_numpy=True)[0].tolist()

# Класс для управления документами
class DocumentManager:
    def __init__(self, data_folder, metadata_file="legal_metadata.json"):
        self.data_folder = data_folder
        self.metadata_file = os.path.join(data_folder, metadata_file)
        self.ensure_metadata_exists()

    def ensure_metadata_exists(self):
        os.makedirs(self.data_folder, exist_ok=True)
        if not os.path.exists(self.metadata_file):
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump({"documents": []}, f, ensure_ascii=False, indent=2)

    def load_metadata(self):
        with open(self.metadata_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def save_metadata(self, metadata):
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

    def calculate_hash(self, file_path):
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def extract_full_text(self, file_path):
        file_ext = os.path.splitext(file_path)[1].lower()
        try:
            if file_ext == '.docx':
                doc = docx.Document(file_path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif file_ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    return "\n".join(page.extract_text() or "" for page in pdf.pages)
            elif file_ext == '.txt':
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    detected_encoding = chardet.detect(raw_data)['encoding']
                with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as file:
                    return file.read()
            return ""
        except Exception as e:
            print(f"[ERROR] Ошибка извлечения текста из {file_path}: {str(e)}")
            return ""

    def get_active_documents(self):
        metadata = self.load_metadata()
        return [doc for doc in metadata.get("documents", []) if doc.get("status") == "active"]

# Функции обработки документов
def split_text_into_sentences(text):
    return [s.text for s in sentenize(text)]

def create_chunks_by_sentence(text, file_metadata, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    sentences = split_text_into_sentences(text)
    chunks = []
    current_chunk = []
    current_size = 0

    for sentence in sentences:
        sentence_size = len(sentence.split())
        if current_size + sentence_size > target_chunk_size and current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(chunk_text)
            if overlap_size > 0:
                overlap_sentences = []
                overlap_tokens = 0
                j = len(current_chunk) - 1
                while j >= 0 and overlap_tokens < overlap_size:
                    overlap_sentences.insert(0, current_chunk[j])
                    overlap_tokens += len(current_chunk[j].split())
                    j -= 1
                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk = [sentence]
                current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size

    if current_chunk:
        chunks.append(" ".join(current_chunk))

    chunk_dicts = []
    for idx, chunk in enumerate(chunks):
        if len(chunk.split()) >= min_chunk_size:
            data = {
                "id": f"{file_metadata.get('file_name', 'unknown')}-chunk-{idx}",
                "text": chunk,
                "metadata": {**file_metadata, "chunk_id": idx, "token_count": len(chunk.split())}
            }
            chunk_dicts.append(data)
    return chunk_dicts

def process_document(file_path, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    file_extension = os.path.splitext(file_path)[1].lower()
    file_metadata = {"file_name": os.path.basename(file_path), "file_path": file_path}
    text = ""

    if file_extension == '.docx':
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        file_metadata["file_type"] = "docx"
    elif file_extension == '.pdf':
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            file_metadata["file_type"] = "pdf"
            file_metadata["num_pages"] = len(pdf.pages)
    elif file_extension == '.txt':
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected_encoding = chardet.detect(raw_data)['encoding']
        with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as file:
            text = file.read()
        file_metadata["file_type"] = "txt"
    else:
        print(f"[INFO] Неподдерживаемый тип файла: {file_path}")
        return []

    if len(text.strip()) < 50:
        print(f"[INFO] Файл {file_metadata['file_name']} содержит менее 50 символов.")
        return []
    return create_chunks_by_sentence(text, file_metadata, target_chunk_size, min_chunk_size, overlap_size)

def process_document_folder(folder_path, **kwargs):
    all_chunks = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')) and not file.startswith('.'):
                file_path = os.path.join(root, file)
                if file != "legal_metadata.json":
                    print(f"[DEBUG] Обработка файла: {file_path}")
                    chunks = process_document(file_path, **kwargs)
                    all_chunks.extend(chunks)
    return all_chunks

# Функции для работы с векторным хранилищем
def create_faiss_index(chunks, embeddings_model):
    texts = [chunk["text"] for chunk in chunks]
    if not texts:
        vector_size = embeddings_model.embed_query("Пустой индекс").shape[0]
        index = faiss.IndexFlatL2(vector_size)
        return index, []

    embed_array = np.array(embeddings_model.embed_documents(texts)).astype("float32")
    vector_size = embed_array.shape[1]
    index = faiss.IndexFlatL2(vector_size)
    index.add(embed_array)
    return index, embed_array

def save_faiss_index(index, metadata, index_path, metadata_path):
    faiss.write_index(index, index_path)
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

def load_or_rebuild_vectorstore(data_folder, indexes_folder, embeddings_model):
    os.makedirs(indexes_folder, exist_ok=True)
    fingerprint_file = os.path.join(indexes_folder, "index_fingerprint.json")
    index_path = os.path.join(indexes_folder, "index.faiss")
    metadata_path = os.path.join(indexes_folder, "document_metadata.json")

    current_fingerprint = {}
    for root, _, files in os.walk(data_folder):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')) and not file.startswith('.'):
                if file != "legal_metadata.json":
                    path = os.path.join(root, file)
                    current_fingerprint[path] = os.path.getmtime(path)

    fingerprint_hash = hashlib.md5(json.dumps(current_fingerprint, sort_keys=True).encode()).hexdigest()

    previous_fingerprint = None
    if os.path.exists(fingerprint_file):
        with open(fingerprint_file, 'r') as f:
            try:
                previous_fingerprint = json.load(f)
            except json.JSONDecodeError:
                previous_fingerprint = None

    if (os.path.exists(index_path) and os.path.exists(metadata_path) and
            previous_fingerprint == fingerprint_hash):
        try:
            print("[INFO] Загрузка существующего векторного хранилища...")
            vectorstore = LC_FAISS.load_local(indexes_folder, embeddings_model, allow_dangerous_deserialization=True)
            return vectorstore
        except Exception as e:
            print(f"[ERROR] Не удалось загрузить векторное хранилище: {e}")
            if os.path.exists(index_path):
                os.remove(index_path)
            if os.path.exists(metadata_path):
                os.remove(metadata_path)

    print("[INFO] Пересоздание векторного индекса...")
    chunks = process_document_folder(data_folder, target_chunk_size=512, min_chunk_size=256, overlap_size=150)

    if not chunks:
        print("[INFO] Не найдено документов для индексации. Создание пустого индекса.")
        vectorstore = LC_FAISS.from_documents(
            [Document(page_content="Пустой индекс", metadata={"source": "empty"})],
            embeddings_model
        )
        vectorstore.save_local(indexes_folder)
        with open(fingerprint_file, 'w') as f:
            json.dump(fingerprint_hash, f)
        return vectorstore

    docs = [Document(page_content=ch["text"], metadata=ch["metadata"]) for ch in chunks]
    vectorstore = LC_FAISS.from_documents(docs, embeddings_model)
    vectorstore.save_local(indexes_folder)
    with open(fingerprint_file, 'w') as f:
        json.dump(fingerprint_hash, f)
    return vectorstore

# Шаблон промпта
def get_legal_prompt_template():
    return (
    "Вы — ThemisBot, экспертный юридический ассистент для жителей. "
    "Отвечайте чётко, уверенно и компетентно, используя только предоставленный контекст. "
    "Никогда не упоминайте, что вы искусственный интеллект или не являетесь юристом. "
    "Отвечайте так, как это сделал бы опытный юридический специалист, дающий практическую помощь. "
    "Давайте ясные, структурированные ответы с маркированными или нумерованными списками при необходимости. "
    "Будьте дружелюбны, лаконичны и действуйте как уверенный профессионал, готовый решить вопрос пользователя. "
    "История чата:\n{chat_history}\n\n"
    "Контекст:\n{context}\n\n"
    "Вопрос: {question}\n\n"
    "Ответ:"
    )

# Класс для управления чатом
class ChatAssistant:
    def __init__(self, qa_chain):
        self.qa = qa_chain
        self.histories = {}

    def get_answer(self, user_query: str, session_id: str = "default"):
        if session_id not in self.histories:
            self.histories[session_id] = []

        chain_history = []
        for i in range(0, len(self.histories[session_id]), 2):
            if i + 1 < len(self.histories[session_id]):
                user_msg = self.histories[session_id][i]["content"]
                assistant_msg = self.histories[session_id][i + 1]["content"]
                chain_history.append((user_msg, assistant_msg))

        try:
            result = self.qa({"question": user_query, "chat_history": chain_history})
            answer = result.get("answer", "")
            if "Это не является юридической консультацией" not in answer:
                answer += "\n\nЭто не является юридической консультацией. Пожалуйста, обратитесь к квалифицированному юристу для профессиональной помощи."
            source_docs = result.get("source_documents", [])
        except Exception as e:
            print(f"[ERROR] Ошибка при обработке запроса: {str(e)}")
            answer = "Извините, произошла ошибка при обработке вашего запроса. Пожалуйста, попробуйте переформулировать вопрос."
            source_docs = []

        self.histories[session_id].append({
            "role": "user",
            "content": user_query,
            "time": datetime.now().strftime("%H:%M:%S")
        })
        self.histories[session_id].append({
            "role": "assistant",
            "content": answer,
            "time": datetime.now().strftime("%H:%M:%S")
        })

        return answer, source_docs

    def clear_history(self, session_id: str = "default"):
        self.histories[session_id] = []

# Инициализация компонентов
embeddings = MyEmbeddings()
doc_manager = DocumentManager(DATA_FOLDER)

try:
    llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, temperature=0, model_name="gpt-4o-mini")
    prompt = PromptTemplate(
        template=get_legal_prompt_template(),
        input_variables=["chat_history", "context", "question"]
    )
    vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER, INDEXES_FOLDER, embeddings)
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 3}),
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": prompt}
    )
    assistant = ChatAssistant(qa_chain)
except Exception as e:
    print(f"[ERROR] Ошибка при инициализации компонентов: {str(e)}")
    raise

# Gradio Interface Functions
def chat_with_bot(user_query, session_id="default"):
    answer, source_docs = assistant.get_answer(user_query, session_id)
    sources = [doc.metadata.get("file_name") for doc in source_docs if doc.metadata.get("file_name")]
    sources = list(set(sources))  # Убираем дубликаты
    return answer, ", ".join(sources) if sources else "Нет источников"

def clear_chat(session_id="default"):
    assistant.clear_history(session_id)
    return "История чата очищена."


# Gradio Interface
css = """
.gradio-container {
    background-color: #1a1a2e !important; 
    max-height: 100vh !important;
    overflow-y: auto;
}

/* Header */
h1, .badge, .header-subtitle {
    color: #e6e6e6 !important;
}

/* Chatbot panel */
#my-chat{
    color: white !important;
    background-color: #16213e !important;
    border: 1px solid #2a3b4d !important;
    border-radius: 8px !important;
    height: 400px !important;
    max-height: 400px !important;
    overflow-y: auto;
}
#my-chat.message {
    color: white !important;
}
/* Message styling */
#my-chat.message.user,
#my-chat.message.bot {
    color: white !important;
    background-color: #0f3460 !important;
    border-radius: 12px !important;
    border: none !important;
    padding: 12px 16px !important;
    box-shadow: none !important;
}

/* Input area */
.input-box {
    background-color: #16213e !important;
    border: 1px solid #2a3b4d !important;
    border-radius: 8px !important;
}

/* Sidebar scroll */
.gr-column:last-child {
    max-height: 500px;
    overflow-y: auto;
}

/* Accordion and sidebar */
.sidebar, .gr-accordion {
    background-color: #16213e !important;
    border: 1px solid #2a3b4d !important;
    border-radius: 8px !important;
}

.gr-accordion .gr-accordion-header {
    background-color: #0f3460 !important;
    color: #ffffff !important;
    border-radius: 6px 6px 0 0 !important;
}

.gr-accordion .gr-accordion-panel {
    background-color: #16213e !important;
    color: #dddddd !important;
    border-radius: 0 0 6px 6px !important;
}

/* Buttons */
button {
    border-radius: 6px !important;
    transition: all 0.3s ease !important;
}

button.primary {
    background-color: #4a6fa5 !important;
    color: white !important;
}

button.primary:hover {
    background-color: #3a5a8a !important;
}

button.secondary {
    background-color: #2a3b4d !important;
    color: white !important;
}

button.secondary:hover {
    background-color: #1a2a3d !important;
}

/* Inputs */
textarea, input[type="text"] {
    background-color: #0f3460 !important;
    color: white !important;
    border: 1px solid #2a3b4d !important;
    border-radius: 6px !important;
    padding: 12px !important;
}

/* Footer text */
.footer-text {
    color: #777777 !important;
    font-size: 0.8em !important;
}
"""


with gr.Blocks(
    theme=gr.themes.Soft(),
    title="ThemisBot - Юридический Ассистент",
    css=css
) as demo:
    with gr.Row():
        with gr.Column(scale=8):
            gr.Markdown(f"""
            <div style="display: flex; align-items: center; gap: 15px; margin-bottom: 10px;">
                <img src="https://i.imgur.com/5L8TbWp.png" width="40"/>
                <div>
                    <h1 style="margin: 0;">ThemisBot</h1>
                    <span class="badge">Юридический ассистент</span>
                </div>
            </div>
            <p class="header-subtitle">
                Получайте быстрые юридические консультации на основе загруженных документов
            </p>
            """)
        with gr.Column(scale=2):
            gr.Markdown(f"""
            <div style="text-align: right; font-size: 0.8em; color: #888;">
                Версия 1.0<br>
                {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
            </div>
            """)

    with gr.Row():
        with gr.Column(scale=3):
            chatbot = gr.Chatbot(
                
                elem_id="my-chat",
                label="Чат с ThemisBot",
                height=400,
                show_copy_button=True,
                show_label=False,
                bubble_full_width=False,
                avatar_images=("🤖", "👤")
            )

            with gr.Row(equal_height=True):
                user_input = gr.Textbox(
                    placeholder="Введите ваш юридический вопрос...",
                    show_label=False,
                    lines=2,
                    max_lines=4,
                    container=False,
                    autofocus=True
                )
                send_button = gr.Button("📨 Отправить", variant="primary", scale=0)

            with gr.Row():
                clear_btn = gr.Button("🧹 Очистить чат", variant="secondary")
                gr.Markdown("<div class='footer-text' style='text-align: right; flex-grow: 1;'>ThemisBot v1.0</div>")

        with gr.Column(scale=1):
            with gr.Accordion("📂 Доступные документы", open=True):
                active_docs = doc_manager.get_active_documents()
                if active_docs:
                    doc_list = "\n".join([f"- {doc['file_name']}" for doc in active_docs[:5]])
                    gr.Markdown(doc_list)
                else:
                    gr.Markdown("Документы не загружены")

            with gr.Accordion("ℹ️ Как пользоваться", open=False):
                gr.Markdown("""
                **Советы для лучших ответов:**
                1. Формулируйте вопросы конкретно  
                2. Указывайте важные детали  
                3. Используйте юридические термины  

                **Примеры вопросов:**  
                - Как расторгнуть договор?  
                - Какие нужны документы?  
                - Какие есть основания?  
                """)

            with gr.Accordion("⚙️ Настройки", open=False):
                gr.Markdown("**Параметры системы:**")
                gr.Slider(minimum=1, maximum=5, value=3, label="Количество источников")
                gr.Checkbox(label="Показывать источники", value=True)


    def chat_response(user_query, chat_history):
        answer, sources = chat_with_bot(user_query, session_id="default")
        if sources:
            answer += f"\n\n🔍 Источники: {sources}"
        chat_history.append((user_query, answer))
        return chat_history, ""

    def clear_chat_history():
        clear_chat(session_id="default")
        return []

 
    send_button.click(
        fn=chat_response,
        inputs=[user_input, chatbot],
        outputs=[chatbot, user_input]
    )

    user_input.submit(
        fn=chat_response,
        inputs=[user_input, chatbot],
        outputs=[chatbot, user_input]
    )

    clear_btn.click(
        fn=clear_chat_history,
        inputs=None,
        outputs=[chatbot]
    )


if __name__ == "__main__":
    demo.launch()
