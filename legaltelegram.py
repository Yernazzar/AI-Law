import os
import json
import hashlib
import shutil
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS as LC_FAISS
from langchain.docstore.document import Document
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
from langchain.prompts.prompt import PromptTemplate
from langchain.embeddings.base import Embeddings
import docx
import pdfplumber
import chardet
from razdel import sentenize
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import warnings
import telebot
import logging
from simulation import SimulationManager, SimulationState
import templates_generation
from typing import List, Optional, Dict, Tuple
from transformers import AutoTokenizer, AutoModel
from telebot import types
from langchain_core.retrievers import BaseRetriever
from langchain_core.documents import Document
import torch
from pydantic import Field
import tax_calculator
from decimal import Decimal
from telebot.handler_backends import State, StatesGroup
from telebot.storage import StateMemoryStorage
import telebot
from telebot import types
import logging
from decimal import Decimal
import tax_calculator
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from io import BytesIO

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    warnings.warn("OPENAI_API_KEY not found in .env file")

# Telegram Token
TELEGRAM_BOT_TOKEN = "token"

# Folders setup
DATA_FOLDER = os.getenv("DATA_FOLDER", "legal_data")
INDEXES_FOLDER = os.getenv("INDEXES_FOLDER", "legal_indexes")
os.makedirs(DATA_FOLDER, exist_ok=True)
os.makedirs(INDEXES_FOLDER, exist_ok=True)

# Global variables
llm = None
vectorstore = None
assistant = None

# Embeddings class for text
class MyEmbeddings(Embeddings):
    def __init__(self):
        self.model = SentenceTransformer('all-MiniLM-L12-v2')

    def embed_documents(self, texts):
        return self.model.encode(texts, convert_to_numpy=True).tolist()

    def embed_query(self, text):
        return self.model.encode([text], convert_to_numpy=True)[0].tolist()

# Document manager class
class DocumentManager:
    def __init__(self, data_folder, metadata_file="legal_metadata.json"):
        self.data_folder = data_folder
        self.metadata_file = os.path.join(data_folder, metadata_file)
        self.ensure_metadata_exists()

    def ensure_metadata_exists(self):
        os.makedirs(self.data_folder, exist_ok=True)
        if not os.path.exists(self.metadata_file):
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump({"documents": []}, f, ensure_ascii=False, indent=2)

    def load_metadata(self):
        with open(self.metadata_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def save_metadata(self, metadata):
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

    def calculate_hash(self, file_path):
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def extract_full_text(self, file_path):
        file_ext = os.path.splitext(file_path)[1].lower()
        try:
            if file_ext == '.docx':
                doc = docx.Document(file_path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif file_ext == '.pdf':
                with pdfplumber.open(file_path) as pdf:
                    return "\n".join(page.extract_text() or "" for page in pdf.pages)
            elif file_ext == '.txt':
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                    detected_encoding = chardet.detect(raw_data)['encoding']
                with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as file:
                    return file.read()
            return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def get_active_documents(self):
        metadata = self.load_metadata()
        return [doc for doc in metadata.get("documents", []) if doc.get("status") == "active"]

    def get_document_path(self, file_name):
        for root, _, files in os.walk(self.data_folder):
            for file in files:
                if file == file_name:
                    return os.path.join(root, file)
        return None

class DocumentReranker:
    def __init__(self, model_name="cross-encoder/ms-marco-MiniLM-L-6-v2"):
        try:
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            self.model = AutoModel.from_pretrained(model_name)
            self.device = "cuda" if torch.cuda.is_available() else "cpu"
            self.model.to(self.device)
            self.model.eval()
            self.is_available = True
            logger.info(f"Document reranker initialized with model: {model_name}")
        except Exception as e:
            logger.error(f"Failed to initialize reranker: {str(e)}")
            self.is_available = False

    def compute_scores(self, query: str, documents: List[Dict]) -> List[Tuple[Dict, float]]:
        if not self.is_available or not documents:
            return [(doc, 1.0) for doc in documents]
        try:
            pairs = [(query, doc["text"]) for doc in documents]
            scores = []
            batch_size = 8
            for i in range(0, len(pairs), batch_size):
                batch_pairs = pairs[i:i + batch_size]
                with torch.no_grad():
                    inputs = self.tokenizer(
                        [pair[0] for pair in batch_pairs],
                        [pair[1] for pair in batch_pairs],
                        padding=True,
                        truncation=True,
                        max_length=512,
                        return_tensors="pt"
                    ).to(self.device)
                    outputs = self.model(**inputs)
                    embeddings = outputs.last_hidden_state[:, 0, :]
                    batch_scores = torch.norm(embeddings, dim=1).cpu().numpy()
                    scores.extend(batch_scores.tolist())
            doc_score_pairs = list(zip(documents, scores))
            doc_score_pairs.sort(key=lambda x: x[1], reverse=True)
            return doc_score_pairs
        except Exception as e:
            logger.error(f"Error in reranker: {str(e)}")
            return [(doc, 1.0) for doc in documents]

    def rerank(self, query: str, documents: List[Dict], top_k: int = None) -> List[Dict]:
        if not documents:
            return []
        doc_score_pairs = self.compute_scores(query, documents)
        if top_k and top_k < len(doc_score_pairs):
            doc_score_pairs = doc_score_pairs[:top_k]
        return [doc for doc, _ in doc_score_pairs]

# Document processing functions
def split_text_into_sentences(text):
    return [s.text for s in sentenize(text)]

def create_chunks_by_sentence(text, file_metadata, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    sentences = split_text_into_sentences(text)
    chunks = []
    current_chunk = []
    current_size = 0
    for sentence in sentences:
        sentence_size = len(sentence.split())
        if current_size + sentence_size > target_chunk_size and current_chunk:
            chunk_text = " ".join(current_chunk)
            chunks.append(chunk_text)
            if overlap_size > 0:
                overlap_sentences = []
                overlap_tokens = 0
                j = len(current_chunk) - 1
                while j >= 0 and overlap_tokens < overlap_size:
                    overlap_sentences.insert(0, current_chunk[j])
                    overlap_tokens += len(current_chunk[j].split())
                    j -= 1
                current_chunk = overlap_sentences + [sentence]
                current_size = sum(len(s.split()) for s in current_chunk)
            else:
                current_chunk = [sentence]
                current_size = sentence_size
        else:
            current_chunk.append(sentence)
            current_size += sentence_size
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    chunk_dicts = []
    for idx, chunk in enumerate(chunks):
        if len(chunk.split()) >= min_chunk_size:
            data = {
                "id": f"{file_metadata.get('file_name', 'unknown')}-chunk-{idx}",
                "text": chunk,
                "metadata": {**file_metadata, "chunk_id": idx, "token_count": len(chunk.split())}
            }
            chunk_dicts.append(data)
    return chunk_dicts

def process_document(file_path, target_chunk_size=512, min_chunk_size=256, overlap_size=50):
    file_extension = os.path.splitext(file_path)[1].lower()
    file_metadata = {"file_name": os.path.basename(file_path), "file_path": file_path}
    text = ""
    if file_extension == '.docx':
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        file_metadata["file_type"] = "docx"
    elif file_extension == '.pdf':
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            file_metadata["file_type"] = "pdf"
            file_metadata["num_pages"] = len(pdf.pages)
    elif file_extension == '.txt':
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            detected_encoding = chardet.detect(raw_data)['encoding']
        with open(file_path, 'r', encoding=detected_encoding or 'utf-8', errors='replace') as file:
            text = file.read()
        file_metadata["file_type"] = "txt"
    else:
        logger.info(f"Unsupported file type: {file_path}")
        return []
    if len(text.strip()) < 50:
        logger.info(f"File {file_metadata['file_name']} contains less than 50 characters.")
        return []
    return create_chunks_by_sentence(text, file_metadata, target_chunk_size, min_chunk_size, overlap_size)

def process_document_folder(folder_path, **kwargs):
    all_chunks = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')) and not file.startswith('.'):
                file_path = os.path.join(root, file)
                if file != "legal_metadata.json":
                    logger.debug(f"Processing file: {file_path}")
                    chunks = process_document(file_path, **kwargs)
                    all_chunks.extend(chunks)
    return all_chunks

# Vector storage functions
def create_faiss_index(chunks, embeddings_model):
    texts = [chunk["text"] for chunk in chunks]
    if not texts:
        vector_size = embeddings_model.embed_query("Empty index").shape[0]
        index = faiss.IndexFlatL2(vector_size)
        return index, []
    embed_array = np.array(embeddings_model.embed_documents(texts)).astype("float32")
    vector_size = embed_array.shape[1]
    index = faiss.IndexFlatL2(vector_size)
    index.add(embed_array)
    return index, embed_array

def save_faiss_index(index, metadata, index_path, metadata_path):
    faiss.write_index(index, index_path)
    with open(metadata_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)

def load_or_rebuild_vectorstore(data_folder, indexes_folder, embeddings_model):
    os.makedirs(indexes_folder, exist_ok=True)
    fingerprint_file = os.path.join(indexes_folder, "index_fingerprint.json")
    index_path = os.path.join(indexes_folder, "index.faiss")
    metadata_path = os.path.join(indexes_folder, "document_metadata.json")
    current_fingerprint = {}
    for root, _, files in os.walk(data_folder):
        for file in files:
            if file.lower().endswith(('.docx', '.pdf', '.txt')) and not file.startswith('.'):
                if file != "legal_metadata.json":
                    path = os.path.join(root, file)
                    current_fingerprint[path] = os.path.getmtime(path)
    fingerprint_hash = hashlib.md5(json.dumps(current_fingerprint, sort_keys=True).encode()).hexdigest()
    previous_fingerprint = None
    if os.path.exists(fingerprint_file):
        with open(fingerprint_file, 'r') as f:
            try:
                previous_fingerprint = json.load(f)
            except json.JSONDecodeError:
                previous_fingerprint = None
    if (os.path.exists(index_path) and os.path.exists(metadata_path) and
            previous_fingerprint == fingerprint_hash):
        try:
            logger.info("Loading existing vector store...")
            vectorstore = LC_FAISS.load_local(indexes_folder, embeddings_model, allow_dangerous_deserialization=True)
            return vectorstore
        except Exception as e:
            logger.error(f"Failed to load vector store: {e}")
            if os.path.exists(index_path):
                os.remove(index_path)
            if os.path.exists(metadata_path):
                os.remove(metadata_path)
    logger.info("Recreating vector index...")
    chunks = process_document_folder(data_folder, target_chunk_size=512, min_chunk_size=256, overlap_size=150)
    if not chunks:
        logger.info("No documents found for indexing. Creating empty index.")
        vectorstore = LC_FAISS.from_documents(
            [Document(page_content="Empty index", metadata={"source": "empty"})],
            embeddings_model
        )
        vectorstore.save_local(indexes_folder)
        with open(fingerprint_file, 'w') as f:
            json.dump(fingerprint_hash, f)
        return vectorstore
    docs = [Document(page_content=ch["text"], metadata=ch["metadata"]) for ch in chunks]
    vectorstore = LC_FAISS.from_documents(docs, embeddings_model)
    vectorstore.save_local(indexes_folder)
    with open(fingerprint_file, 'w') as f:
        json.dump(fingerprint_hash, f)
    return vectorstore

# Prompt template
def get_legal_prompt_template():
    return (
        "Роль: Ты — ThemisBot, юридический помощник для жителей Казахстана. Твоя задача — давать точные, понятные и уверенные ответы по юридическим вопросам, основываясь на законах Казахстана и документах из базы. Ты также проводишь симуляции, чтобы проверить знания пользователей в игровой форме, выступая как опытный юрист."
        "Используй исключительно законы и правила Казахстана. Ссылайся на документы из базы, если они есть. Если данных нет, отвечай на основе общих знаний, оставаясь уверенным."
        "Отвечайте на языке запроса. "
        "Внимательно анализируйте вопрос и контекст перед ответом. "
        "Будь разговорным, дружелюбным и уверенным, как будто объясняешь другу. Избегай сложных юридических терминов, а если они нужны — объясняй их простыми словами."
        "Всегда используй маркированные списки (буллет-поинты) для ответов с несколькими пунктами."
        "Упоминай документы, если они использованы (например, «По данным [Название документа]»)."
        "Если вопрос непонятен, попроси уточнить: «Можешь рассказать подробнее, что ты имеешь в виду?»"
        "Если информации нет, дай общий ответ: «Точных данных по этому вопросу у меня нет, но в целом ситуация выглядит так...» и предложи возможное решение."
        "Не говори, что ты не юрист или что твои ответы не являются юридической консультацией."
        "Не ссылайся на законы других стран, если пользователь явно не просит."
        "Будьте дружелюбны и выступайте в роли полезного гида. "
        "Думайте и отвечайте как юрист, будьте уверенны в своих ответах"
        "Не говори -  Это не является юридической консультацией. Пожалуйста, обратитесь к квалифицированному юристу для профессиональной помощи. Будь уверен в своих ответах"
        "История чата:\n{chat_history}\n\n"
        "Контекст:\n{context}\n\n"
        "Вопрос: {question}\n\n"
        "Ответ:"
    )

class RerankedRetriever(BaseRetriever):
    """
    Custom retriever that uses vector search followed by reranking.
    """
    vectorstore_retriever: BaseRetriever = Field(description="Base vector store retriever")
    reranker: "DocumentReranker" = Field(description="Document reranker instance")
    k: int = Field(default=3, description="Number of final documents to retrieve")
    rerank_top_k: int = Field(default=10, description="Number of initial documents to consider for reranking")

    def _get_relevant_documents(self, query: str, *, run_manager=None) -> List[Document]:
        try:
            logger.info(f"Retrieving documents for query: {query}")
            initial_docs = self.vectorstore_retriever.get_relevant_documents(query)
            if not initial_docs or not self.reranker.is_available:
                logger.info("No initial docs or reranker unavailable, returning top k")
                return initial_docs[:self.k]
            doc_dicts = [
                {"text": doc.page_content, "metadata": doc.metadata, "doc_obj": doc}
                for doc in initial_docs
            ]
            reranked_docs = self.reranker.rerank(query, doc_dicts, top_k=self.k)
            return [doc["doc_obj"] for doc in reranked_docs]
        except Exception as e:
            logger.error(f"Error in reranked retrieval: {str(e)}")
            return self.vectorstore_retriever.get_relevant_documents(query)[:self.k]

def setup_qa_chain_with_reranking(llm, vectorstore, rerank_enabled=True):
    logger.info("Setting up QA chain with reranking...")
    prompt = PromptTemplate(
        template=get_legal_prompt_template(),
        input_variables=["chat_history", "context", "question"]
    )
    base_retriever = vectorstore.as_retriever(search_kwargs={"k": 10})
    logger.info(f"Base retriever created: {base_retriever}")
    if rerank_enabled:
        try:
            reranker = DocumentReranker()
            logger.info(f"Reranker initialized: {reranker.is_available}, type: {type(reranker)}")
            if reranker.is_available:
                logger.info(f"Creating RerankedRetriever with vectorstore_retriever={base_retriever}, reranker={reranker}")
                retriever = RerankedRetriever(
                    vectorstore_retriever=base_retriever,
                    reranker=reranker,
                    k=3,
                    rerank_top_k=10
                )
                logger.info("RerankedRetriever created successfully")
            else:
                retriever = base_retriever
                logger.warning("Reranker not available, falling back to vector search only.")
        except Exception as e:
            logger.error(f"Error setting up reranker: {str(e)}")
            retriever = base_retriever
    else:
        retriever = base_retriever
    logger.info("Creating ConversationalRetrievalChain...")
    qa_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever=retriever,
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": prompt}
    )
    logger.info("QA chain created successfully")
    return qa_chain

# Chat management class
class ChatAssistant:
    def __init__(self, qa_chain):
        self.qa = qa_chain
        self.histories = {}

    def get_answer(self, user_query: str, session_id: str = "default"):
        if session_id not in self.histories:
            self.histories[session_id] = []
        chain_history = []
        for i in range(0, len(self.histories[session_id]), 2):
            if i + 1 < len(self.histories[session_id]):
                user_msg = self.histories[session_id][i]["content"]
                assistant_msg = self.histories[session_id][i + 1]["content"]
                chain_history.append((user_msg, assistant_msg))
        try:
            result = self.qa({"question": user_query, "chat_history": chain_history})
            answer = result.get("answer", "")
            if "Это не является юридической консультацией" not in answer:
                answer += "\n\nЭто не является юридической консультацией. Пожалуйста, обратитесь к квалифицированному юристу для профессиональной помощи."
            source_docs = result.get("source_documents", [])
        except Exception as e:
            logger.error(f"Error processing request: {str(e)}")
            answer = "Извините, произошла ошибка при обработке вашего запроса. Пожалуйста, попробуйте переформулировать вопрос."
            source_docs = []
        self.histories[session_id].append({
            "role": "user",
            "content": user_query,
            "time": datetime.now().strftime("%H:%M:%S")
        })
        self.histories[session_id].append({
            "role": "assistant",
            "content": answer,
            "time": datetime.now().strftime("%H:%M:%S")
        })
        return answer, source_docs

    def clear_history(self, session_id: str = "default"):
        self.histories[session_id] = []

llm = None
vectorstore = None
assistant = None

# Initialize components
logger.info("Initializing components...")
embeddings = MyEmbeddings()
doc_manager = DocumentManager(DATA_FOLDER)
reranking_status = {}

try:
    logger.info("Setting up LLM and vector store...")
    llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, temperature=0, model_name="gpt-4o-mini")
    vectorstore = load_or_rebuild_vectorstore(DATA_FOLDER, INDEXES_FOLDER, embeddings)
    prompt = PromptTemplate(
        template=get_legal_prompt_template(),
        input_variables=["chat_history", "context", "question"]
    )
    enable_reranking = os.getenv("ENABLE_RERANKING", "True").lower() == "true"
    qa_chain = setup_qa_chain_with_reranking(llm, vectorstore, rerank_enabled=enable_reranking)
    assistant = ChatAssistant(qa_chain)
    logger.info(f"Components initialized successfully. Reranking: {'enabled' if enable_reranking else 'disabled'}")
    logger.info("Initializing SimulationManager...")
    simulation_manager = SimulationManager(qa_chain, vectorstore)
    logger.info("Simulation manager initialized successfully.")
except Exception as e:
    logger.error(f"Error initializing components: {str(e)}")
    raise

# Initialize Telegram bot
logger.info("Starting Telegram bot...")

state_storage = StateMemoryStorage()
bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN, state_storage=state_storage)

# Command handlers
@bot.message_handler(commands=['simulation'])
def handle_simulation(message):
    user_id = str(message.from_user.id)
    command_args = message.text.split(' ', 1)
    topic = command_args[1] if len(command_args) > 1 else None
    bot.send_chat_action(message.chat.id, 'typing')
    response = simulation_manager.start_simulation(user_id, topic)
    if len(response) > 4000:
        chunks = [response[i:i + 4000] for i in range(0, len(response), 4000)]
        for chunk in chunks:
            bot.send_message(message.chat.id, chunk, parse_mode='Markdown')
    else:
        bot.send_message(message.chat.id, response, parse_mode='Markdown')

@bot.message_handler(commands=['stop_simulation'])
def handle_stop_simulation(message):
    user_id = str(message.from_user.id)
    if simulation_manager.is_in_simulation(user_id):
        bot.send_chat_action(message.chat.id, 'typing')
        response = simulation_manager.end_simulation(user_id)
        if len(response) > 4000:
            chunks = [response[i:i + 4000] for i in range(0, len(response), 4000)]
            for chunk in chunks:
                bot.send_message(message.chat.id, chunk, parse_mode='Markdown')
        else:
            bot.send_message(message.chat.id, response, parse_mode='Markdown')
    else:
        bot.send_message(message.chat.id, "У вас нет активной симуляции.")

bot.user_data = {}

@bot.message_handler(commands=['taxcalc'])
def handle_taxcalc(message):
    """Обработчик команды /taxcalc <сумма>."""
    try:
        # Извлечение суммы из команды
        args = message.text.split()
        if len(args) != 2:
            bot.reply_to(message, "Используйте: /taxcalc <сумма> (например, /taxcalc 500000)")
            logging.warning(f"User {message.from_user.id} used invalid taxcalc format: {message.text}")
            return
        
        amount = args[1]
        valid, error, value = tax_calculator.validate_input(amount, "decimal")
        if not valid:
            bot.reply_to(message, error)
            logging.warning(f"User {message.from_user.id} entered invalid amount: {amount}")
            return
        
        # Сохранение суммы в user_data
        bot.user_data[message.from_user.id] = {"amount": value}
        
        # Предложение выбора формы расчета
        keyboard = types.InlineKeyboardMarkup()
        keyboard.add(types.InlineKeyboardButton("Зарплата на руки", callback_data=f"tax:salary_net:{value}"))
        keyboard.add(types.InlineKeyboardButton("Оклад от суммы на руки", callback_data=f"tax:salary_gross:{value}"))
        keyboard.add(types.InlineKeyboardButton("ИП на упрощенке", callback_data=f"tax:ip_uproschenka:{value}"))
        
        bot.reply_to(message, f"Сумма: {value:,.2f} тенге. Выберите форму расчета:", reply_markup=keyboard)
        logging.info(f"User {message.from_user.id} started taxcalc with amount: {value}")
    
    except Exception as e:
        bot.reply_to(message, "Ошибка при обработке команды. Попробуйте снова.")
        logging.error(f"User {message.from_user.id} error in taxcalc: {str(e)}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("tax:"))
def handle_tax_selection(call):
    """Обработчик выбора формы расчета."""
    try:
        bot.answer_callback_query(call.id)
        parts = call.data.split(":")
        tax_type = parts[1]
        amount = Decimal(parts[2])
        
        # Сохранение типа расчета
        user_data = bot.user_data.get(call.from_user.id, {})
        user_data["tax_type"] = tax_type
        user_data["amount"] = amount
        bot.user_data[call.from_user.id] = user_data
        
        # Запрос резидентства
        keyboard = types.InlineKeyboardMarkup()
        keyboard.add(types.InlineKeyboardButton("Да", callback_data="resident:yes"))
        keyboard.add(types.InlineKeyboardButton("Нет", callback_data="resident:no"))
        bot.send_message(call.message.chat.id, "Вы резидент РК?", reply_markup=keyboard)
        logging.info(f"User {call.from_user.id} selected tax type: {tax_type}, amount: {amount}")
    
    except Exception as e:
        bot.send_message(call.message.chat.id, "Ошибка при выборе формы. Начните заново с /taxcalc.")
        logging.error(f"User {call.from_user.id} error in tax_selection: {str(e)}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("resident:"))
def handle_resident_selection(call):
    """Обработчик выбора статуса резидента."""
    try:
        bot.answer_callback_query(call.id)
        is_resident = call.data == "resident:yes"
        
        # Сохранение резидентства
        user_data = bot.user_data.get(call.from_user.id, {})
        user_data["is_resident"] = is_resident
        bot.user_data[call.from_user.id] = user_data
        
        # Запрос вычета 14 МРП
        keyboard = types.InlineKeyboardMarkup()
        keyboard.add(types.InlineKeyboardButton("Да", callback_data="deduction:yes"))
        keyboard.add(types.InlineKeyboardButton("Нет", callback_data="deduction:no"))
        bot.send_message(call.message.chat.id, "Применять вычет 14 МРП?", reply_markup=keyboard)
        logging.info(f"User {call.from_user.id} selected resident: {is_resident}")
    
    except Exception as e:
        bot.send_message(call.message.chat.id, "Ошибка при выборе резидентства. Начните заново с /taxcalc.")
        logging.error(f"User {call.from_user.id} error in resident_selection: {str(e)}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("deduction:"))
def handle_deduction_selection(call):
    """Обработчик выбора вычета 14 МРП."""
    try:
        bot.answer_callback_query(call.id)
        has_deduction = call.data == "deduction:yes"
        
        # Сохранение вычета
        user_data = bot.user_data.get(call.from_user.id, {})
        user_data["has_deduction"] = has_deduction
        bot.user_data[call.from_user.id] = user_data
        
        # Запрос статуса инвалида
        keyboard = types.InlineKeyboardMarkup()
        keyboard.add(types.InlineKeyboardButton("Да", callback_data="disabled:yes"))
        keyboard.add(types.InlineKeyboardButton("Нет", callback_data="disabled:no"))
        bot.send_message(call.message.chat.id, "Вы инвалид?", reply_markup=keyboard)
        logging.info(f"User {call.from_user.id} selected deduction: {has_deduction}")
    
    except Exception as e:
        bot.send_message(call.message.chat.id, "Ошибка при выборе вычета. Начните заново с /taxcalc.")
        logging.error(f"User {call.from_user.id} error in deduction_selection: {str(e)}")

@bot.callback_query_handler(func=lambda call: call.data.startswith("disabled:"))
def handle_disabled_selection(call):
    """Обработчик выбора статуса инвалида и финальный расчет."""
    try:
        bot.answer_callback_query(call.id)
        is_disabled = call.data == "disabled:yes"
        
        # Получение данных
        user_data = bot.user_data.get(call.from_user.id, {})
        tax_type = user_data.get("tax_type")
        amount = user_data.get("amount")
        params = {
            "is_resident": user_data.get("is_resident", True),
            "has_deduction": user_data.get("has_deduction", True),
            "is_disabled": is_disabled
        }
        
        # Выполнение расчета
        if tax_type == "salary_net":
            params["gross"] = amount
            tax, description = tax_calculator.calculate_tax("salary_net", params)
            result = f"Расчет зарплаты:\n{description}"
        elif tax_type == "salary_gross":
            params["net"] = amount
            tax, description = tax_calculator.calculate_tax("salary_gross", params)
            result = f"Расчет оклада:\n{description}"
        elif tax_type == "ip_uproschenka":
            params["income"] = amount
            tax, description = tax_calculator.calculate_tax("ip_uproschenka", params)
            result = f"Расчет для ИП на упрощенке:\n{description}"
        else:
            result = "Ошибка: неизвестный тип расчета."
        
        # Отправка результата
        bot.send_message(call.message.chat.id, result)
        
        # Очистка данных
        bot.user_data.pop(call.from_user.id, None)
        logging.info(f"User {call.from_user.id} calculated {tax_type}: {tax}")
    
    except Exception as e:
        logging.error(f"User {call.from_user.id} error in disabled_selection: {str(e)}")


@bot.message_handler(commands=['template'])
def handle_template(message):
    """Обработчик команды /template для поиска и отправки шаблонов документов."""
    # Извлечение запроса после команды
    user_query = message.text.replace('/template', '').strip().lower()
    
    if not user_query:
        bot.reply_to(message, "Пожалуйста, укажите запрос после команды, например: /template документы для развода")
        return
    
    # Используем функцию из templates_generation.py
    response = templates_generation.generate_template_response(user_query)
    
    if response["success"]:
        reply_text = response["message"] + "\n\n" + "\n".join(
            f"- {tpl['name']}" for tpl in response["templates"]
        )
        bot.reply_to(message, reply_text)
        
        # Отправка файлов шаблонов, если они найдены
        for tpl in response["templates"]:
            if tpl.get("file") and not tpl.get("error"):
                try:
                    with open(tpl["file"], "rb") as file:
                        bot.send_document(message.chat.id, document=file, caption=tpl["name"])
                except Exception as e:
                    bot.reply_to(message, f"Ошибка при отправке файла {tpl['name']}: {str(e)}")
            elif tpl.get("error"):
                bot.reply_to(message, tpl["error"])
    else:
        bot.reply_to(message, response["message"])

@bot.message_handler(commands=['start'])
def handle_start(message):
    user_id = str(message.from_user.id)
    bot.send_message(
        message.chat.id,
        "👋 Здравствуйте! Я ThemisBot - ваш юридический ассистент.\n\n"
        "Задавайте мне юридические вопросы, и я постараюсь на них ответить.\n\n"
        "Команды:\n"
        "/help - Показать справку\n"
        "/clear - Очистить историю нашего разговора\n"
        "/template [тема] - Запросить шаблон документов\n"
        "/simulation [тема] - Начать юридическую симуляцию\n"
        "/stop_simulation - Остановить текущую симуляцию\n"
        "/taxcalc [сумма]- Калькулятор налогов\n"
        "/toggle_reranking - Включить/выключить улучшенный поиск документов"
    )
    #assistant.clear_history(user_id)

@bot.message_handler(commands=['help'])
def handle_help(message):
    bot.send_message(
        message.chat.id,
        "Я - ваш юридический ассистент. Вот что я могу:\n\n"
        "1. *Консультация:* Задавайте юридические вопросы, и я отвечу на основе базы знаний\n"
        "2. *Симуляция:* Используйте /simulation [тема], чтобы начать симуляцию юридической ситуации\n\n"
        "Доступные команды:\n"
        "- /start - Начать диалог\n"
        "- /help - Показать эту справку\n"
        "- /clear - Очистить историю разговора\n"
        "- /template [тема] - Запросить шаблон документов\n"
        "- /simulation [тема] - Начать юридическую симуляцию\n"
        "- /stop_simulation - Остановить текущую симуляцию\n"
        "- /taxcalc [сумма]- Калькулятор налогов\n"
        "- /toggle_reranking - Включить/выключить улучшенный поиск документов\n",
        parse_mode='Markdown'
    )

@bot.message_handler(commands=['clear'])
def handle_clear(message):
    user_id = str(message.from_user.id)
    assistant.clear_history(user_id)
    bot.send_message(
        message.chat.id,
        "🗑️ История разговора очищена. Вы можете задать новый вопрос."
    )

@bot.message_handler(commands=['toggle_reranking'])
def handle_toggle_reranking(message):
    user_id = str(message.from_user.id)
    current_status = reranking_status.get(user_id, True)
    new_status = not current_status
    reranking_status[user_id] = new_status
    global vectorstore, llm
    prompt = PromptTemplate(
        template=get_legal_prompt_template(),
        input_variables=["chat_history", "context", "question"]
    )
    base_retriever = vectorstore.as_retriever(search_kwargs={"k": 3 if not new_status else 10})
    if new_status:
        try:
            reranker = DocumentReranker()
            if reranker.is_available:
                retriever = RerankedRetriever(
                    vectorstore_retriever=base_retriever,
                    reranker=reranker,
                    k=3,
                    rerank_top_k=10
                )
                status_message = "✅ Reranking включен. Результаты будут отсортированы по релевантности."
            else:
                retriever = base_retriever
                status_message = "⚠️ Не удалось инициализировать reranking. Используется обычный поиск."
        except Exception as e:
            logger.error(f"Error toggling reranking: {str(e)}")
            retriever = base_retriever
            status_message = "⚠️ Ошибка при включении reranking. Используется обычный поиск."
    else:
        retriever = base_retriever
        status_message = "❌ Reranking отключен. Используется векторный поиск."
    user_qa_chain = ConversationalRetrievalChain.from_llm(
        llm,
        retriever=retriever,
        return_source_documents=True,
        combine_docs_chain_kwargs={"prompt": prompt}
    )
    assistant.qa = user_qa_chain
    bot.send_message(message.chat.id, status_message)

@bot.message_handler(func=lambda message: True)
def handle_message(message):
    user_id = str(message.from_user.id)
    user_query = message.text
    bot.send_chat_action(message.chat.id, 'typing')

    if simulation_manager.is_in_simulation(user_id):
        sim_state = simulation_manager.get_simulation_state(user_id)
        if sim_state == SimulationState.SETUP:
            response = simulation_manager.start_simulation(user_id, user_query)
        elif sim_state == SimulationState.RUNNING:
            response = simulation_manager.process_answer(user_id, user_query)
        else:
            response = simulation_manager.end_simulation(user_id)
        if len(response) > 4000:
            chunks = [response[i:i + 4000] for i in range(0, len(response), 4000)]
            for chunk in chunks:
                bot.send_message(message.chat.id, chunk, parse_mode='Markdown')
        else:
            bot.send_message(message.chat.id, response, parse_mode='Markdown')
    else:
        answer, source_docs = assistant.get_answer(user_query, user_id)
        response_message = answer
        if len(response_message) > 4000:
            chunks = [response_message[i:i + 4000] for i in range(0, len(response_message), 4000)]
            for chunk in chunks:
                bot.send_message(message.chat.id, chunk, parse_mode='Markdown')
        else:
            bot.send_message(message.chat.id, response_message, parse_mode='Markdown')
        source_files = {}
        for doc in source_docs:
            if "file_name" in doc.metadata and "file_path" in doc.metadata:
                file_name = doc.metadata["file_name"]
                file_path = doc.metadata["file_path"]
                source_files[file_name] = file_path
        if source_files:
            bot.send_message(
                message.chat.id,
                "📚 *Использованные документы:*",
                parse_mode='Markdown'
            )
            for file_name, file_path in source_files.items():
                try:
                    if os.path.exists(file_path):
                        with open(file_path, 'rb') as file:
                            bot.send_document(
                                message.chat.id,
                                file,
                                caption=f"Документ: {file_name}"
                            )
                    else:
                        actual_path = doc_manager.get_document_path(file_name)
                        if actual_path and os.path.exists(actual_path):
                            with open(actual_path, 'rb') as file:
                                bot.send_document(
                                    message.chat.id,
                                    file,
                                    caption=f"Документ: {file_name}"
                                )
                except Exception as e:
                    logger.error(f"Error sending document {file_name}: {e}")
                    bot.send_message(
                        message.chat.id,
                        f"Не удалось отправить документ: {file_name}"
                    )

if __name__ == "__main__":
    logger.info("Starting Themis Telegram bot...")
    try:
        bot.polling(none_stop=True)
    except Exception as e:
        logger.error(f"Error in bot polling: {e}")